import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

class configureReader:
    def __init__(self):
        self.comp = "Название компании"
        self.director = "ФИО"
        self.city = "Город"
        self.num = 0
        self.path = os.getcwd() + "/config.txt"
        if not self.read_file():
            self.write_file()

    def read_file(self):
        if not os.path.exists(self.path):
            return False
        with open("config.txt", "r") as f:
            try:
                lines = f.readlines()
                self.comp = lines[0]
                self.director = lines[1]
                self.city = lines[2]
                self.num = int(lines[3])
                return True
            except:
                self.comp = "Название компании"
                self.director = "ФИО"
                self.city = "Город"
                self.num = 0
        return False

    def write_file(self):
        with open("config.txt", "w") as f:
            print(self.comp)
            f.writelines([str(self.comp).replace("\n", "") + "\n", str(self.director).replace("\n", "") + "\n", str(self.city).replace("\n", "") + "\n", str(self.num).replace("\n", "") + "\n"])


class docWriter:
    def __init__(self):
        self.docReset()
        self.c = configureReader()

    def docReset(self):
        self.doc = Document()
        self.style = self.doc.styles['Normal']
        self.font = self.style.font
        self.font.name = 'Times New Roman'
        self.font.size = Pt(12)
        self.path = os.getcwd() + "/Документы/"
        self.path_Dismissal = self.path + "/Приказы об увольнении/"
        self.path_Hiring = self.path + "/Приказы о зачислении/"
        if not os.path.exists(self.path):
            os.makedirs(self.path)
        if not os.path.exists(self.path_Dismissal):
            os.makedirs(self.path_Dismissal)
        if not os.path.exists(self.path_Hiring):
            os.makedirs(self.path_Hiring)

    def docDismissal(self, last_name, first_name, middle_name, date, position):
        self.docReset()
        p1 = self.doc.add_paragraph()
        p1.add_run(self.c.comp)
        p1.runs[0].font.size = Pt(14)
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = self.doc.add_paragraph()
        p2.add_run("Приказ об увольнении").bold = True
        p2.runs[0].font.size = Pt(18)
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3 = self.doc.add_paragraph()
        p3.add_run(self.c.city)
        p4 = self.doc.add_paragraph("№")
        p4.add_run(str(self.c.num))
        self.c.num += 1
        p4.add_run(" от " + date)
        p5 = self.doc.add_paragraph("На основании заявления об увольнении")
        p5.add_run(" причина ").italic = True
        p5.add_run(position + " " + last_name + " " + first_name + " " + middle_name)
        p5.add_run(" от ")
        p5.add_run(" дата ").italic = True
        p6 = self.doc.add_paragraph("ПРИКАЗЫВАЮ:")
        p7 = self.doc.add_paragraph("1. Уволить " + date + " " + position + " " + last_name + " " + first_name + " " + middle_name)
        p7.add_run(" причина , статья трудового кодекса").italic = True
        p8 = self.doc.add_paragraph("2. Специалисту отдела кадров ")
        p8.add_run(" ФИО ").italic = True
        p8.add_run("расторгнуть трудовой договор " + date + " №")
        p8.add_run("Номер").italic = True
        p8.add_run(" от ")
        p8.add_run("дата трудового договора, ").italic = True
        p8.add_run("заключенный с " + last_name + " " + first_name + " " + middle_name + " и внести запись в трудовую книжку.")
        p9 = self.doc.add_paragraph("3. Бухгалтеру ")
        p9.add_run(" ФИО ").italic = True
        p9.add_run("произвести полный расчет " + position + " " + last_name + " " + first_name + " " + middle_name + ".")
        p9 = self.doc.add_paragraph("4. Контроль за исполнением приказа возлагаю на руководителя отдела кадров ")
        p9.add_run(" ФИО ").italic = True
        p10 = self.doc.add_paragraph("Генеральный директор ")
        if self.c.director == "ФИО":
            p10.add_run(" Фамилия ").italic = True
        else:
            p10.add_run(self.c.director.split()[0])
        p10.add_run(self.c.director)
        p11 = self.doc.add_paragraph("С приказом ознакомлены:")
        p1.style = self.doc.styles['Normal']
        p2.style = self.doc.styles['Normal']
        p3.style = self.doc.styles['Normal']
        p4.style = self.doc.styles['Normal']
        p5.style = self.doc.styles['Normal']
        p6.style = self.doc.styles['Normal']
        p7.style = self.doc.styles['Normal']
        p8.style = self.doc.styles['Normal']
        p9.style = self.doc.styles['Normal']
        p10.style = self.doc.styles['Normal']
        p11.style = self.doc.styles['Normal']

        if not os.path.exists(self.path_Dismissal + "Увольнение " + str(last_name) + str(first_name) + str(middle_name) + ".docx"):
            self.doc.save(self.path_Dismissal + "Увольнение " + str(last_name) + str(first_name) + str(middle_name) + ".docx")
            self.c.write_file()
            return True
        return False

    def docHiring(self, last_name, first_name, middle_name, date, position, salary, office):
        self.docReset()
        p1 = self.doc.add_paragraph("Форма приказа о приеме на работу")
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2 = self.doc.add_paragraph("утверждена приказом директора ")
        p2.add_run(self.c.comp)
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3 = self.doc.add_paragraph(" от ")
        p3.add_run(" дата ").italic = True
        p3.add_run("№ ")
        p3.add_run(str(self.c.num))
        p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p4 = self.doc.add_paragraph()
        p4.add_run(self.c.comp)
        p4.runs[0].font.size = Pt(14)
        p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5 = self.doc.add_paragraph()
        p5.add_run("Приказ").bold = True
        p5.runs[0].font.size = Pt(18)
        p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p6 = self.doc.add_paragraph()
        p6.add_run(self.c.city)
        p7 = self.doc.add_paragraph("№")
        p7.add_run(str(self.c.num))
        self.c.num += 1
        p7.add_run(" от " + date)
        p8 = self.doc.add_paragraph("О приеме на работу " + last_name + " " + first_name + " " + middle_name + ".")
        p8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p9 = self.doc.add_paragraph("Принять " + last_name + " " + first_name + " " + middle_name + " с " + date + " в " + office + "на должность " + position + " с окладом " + salary + " руб.")
        p10 = self.doc.add_paragraph("Основание: трудовой договор от " + date + " № ")
        p10.add_run("Номер").italic = True
        p10.add_run(".")
        p11 = self.doc.add_paragraph("Генеральный директор ")
        if self.c.director == "ФИО":
            p11.add_run(" Фамилия ").italic = True
        else:
            p11.add_run(self.c.director.split()[0])
        p11.add_run(self.c.director)
        p12 = self.doc.add_paragraph("С приказом ознакомлены:")
        p1.style = self.doc.styles['Normal']
        p2.style = self.doc.styles['Normal']
        p3.style = self.doc.styles['Normal']
        p4.style = self.doc.styles['Normal']
        p5.style = self.doc.styles['Normal']
        p6.style = self.doc.styles['Normal']
        p7.style = self.doc.styles['Normal']
        p8.style = self.doc.styles['Normal']
        p9.style = self.doc.styles['Normal']
        p10.style = self.doc.styles['Normal']
        p11.style = self.doc.styles['Normal']
        p12.style = self.doc.styles['Normal']

        if not os.path.exists(self.path_Hiring + "Зачисление " + str(last_name) + str(first_name) + str(middle_name) + ".docx"):
            self.doc.save(self.path_Hiring + "Зачисление " + str(last_name) + str(first_name) + str(middle_name) + ".docx")
            self.c.write_file()
            return True
        return False
